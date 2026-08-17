"""Small, bounded extractors for user-uploaded text documents."""

from __future__ import annotations

from io import BytesIO
from pathlib import PurePath

from app.media import normalize_text

MAX_DOCUMENT_PAGES = 100
MAX_EXTRACTED_CHARACTERS = 10_000


def extract_document(filename: str, content_type: str, content: bytes) -> str:
    suffix = PurePath(filename).suffix.lower()
    if suffix == ".txt" or content_type == "text/plain":
        text = _extract_text(content)
    elif suffix == ".pdf" or content_type == "application/pdf":
        text = _extract_pdf(content)
    elif suffix == ".docx" or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text = _extract_docx(content)
    else:
        raise ValueError("unsupported document type")
    text = normalize_text(text)
    if not text:
        raise ValueError("document contains no extractable text")
    if len(text) > MAX_EXTRACTED_CHARACTERS:
        raise ValueError("document exceeds the 10,000 character limit")
    return text


def _extract_text(content: bytes) -> str:
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError as error:
        raise ValueError("text files must use UTF-8 encoding") from error


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))
    if reader.is_encrypted:
        raise ValueError("encrypted PDFs are not supported")
    if len(reader.pages) > MAX_DOCUMENT_PAGES:
        raise ValueError("PDF exceeds the 100 page limit")
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(content: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        paragraphs.extend(cell.text for row in table.rows for cell in row.cells)
    return "\n\n".join(paragraphs)
